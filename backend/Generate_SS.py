from docx import Document
import re

def generate_salary_slip(template_path, employee_data, headers):
    
    # Load the Word template
    template = Document(template_path)
    
    # Map employee data to headers (placeholders)
    placeholder_data = dict(zip(headers, employee_data))
    
    # Replace placeholders in the document
    for paragraph in template.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in placeholder_data.items():
                if f"{{{placeholder}}}" in run.text:
                    run.text = run.text.replace(f"{{{placeholder}}}", value)
    
    # Replace placeholders in tables
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in placeholder_data.items():
                            if f"{{{placeholder}}}" in run.text:
                                run.text = run.text.replace(f"{{{placeholder}}}", value)
    
    return template
